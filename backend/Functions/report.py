# Functions/report.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, Cm
import google.generativeai as genai
import base64
from io import BytesIO
from PIL import Image
import tempfile
import os
import json
import re
from datetime import datetime
from typing import List, Dict


class ReportGenerator:
    def __init__(self):
        self.model = genai.GenerativeModel('gemini-2.0-flash-thinking-exp')

    def format_text(self, text: str) -> str:
        """Format markdown-style text to plain text"""
        # Remove markdown-style formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
        text = re.sub(r'#{1,6}\s+', '', text)  # Remove headers
        return text.strip()

    def add_formatted_paragraph(self, doc: Document, text: str, style: str = 'Normal') -> None:
        """Add a formatted paragraph to the document"""
        # Split the text into segments based on formatting
        segments = re.split(r'(\*\*.*?\*\*)', text)

        paragraph = doc.add_paragraph(style=style)
        for segment in segments:
            if segment.startswith('**') and segment.endswith('**'):
                # Bold text
                run = paragraph.add_run(segment.strip('*'))
                run.bold = True
            else:
                # Normal text
                paragraph.add_run(segment)

    def setup_document_styles(self, doc: Document):
        """Setup document styles"""
        # Heading 1
        h1_style = doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 51, 102)
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)

        # Heading 2
        h2_style = doc.styles['Heading 2']
        h2_font = h2_style.font
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(6)

        # Normal text
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(12)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # List style
        list_style = doc.styles.add_style('List Item', WD_STYLE_TYPE.PARAGRAPH)
        list_font = list_style.font
        list_font.size = Pt(11)
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)

    def generate_summary(self, analysis_data: Dict, anomaly_data: Dict) -> str:
        """Generate executive summary"""
        prompt = f"""
        Create a professional financial analysis report with these sections:

        1. Executive Summary
        - Overview of findings
        - Financial highlights
        - Critical insights

        2. Risk Areas
        Based on anomaly data:
        {json.dumps(anomaly_data, indent=2)}
        - Key risk factors identified
        - Potential impacts
        - Areas of concern

        3. Recommendations
        - Action items (prioritized)
        - Improvement areas
        - Strategic initiatives

        Format in clear language with bullet points. Be specific and data-driven.
        Highlight important points using ** for emphasis (e.g., **Critical Risk**).
        """

        response = self.model.generate_content(prompt)
        return response.text

    def add_image_with_caption(self, doc: Document, image_base64: str, caption: str):
        """Add image with caption"""
        try:
            image_binary = base64.b64decode(image_base64)
            image_stream = BytesIO(image_binary)
            image = Image.open(image_stream)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                image.save(tmp.name)
                doc.add_picture(tmp.name, width=Inches(6.0))

                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(caption)
                caption_run.bold = True
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()
                os.unlink(tmp.name)
        except Exception as e:
            print(f"Failed to add image: {str(e)}")

    def create_word_report(
            self,
            analysis_data: Dict,
            anomaly_data: Dict,
            file_names: List[str],
            conversation_id: str
    ) -> str:
        doc = Document()
        self.setup_document_styles(doc)

        # Title Page
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Financial Analysis Report')
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run('CONFIDENTIAL\n').bold = True
        meta_para.add_run(f'\nGenerated: {datetime.now().strftime("%B %d, %Y %H:%M:%S")}\n')
        meta_para.add_run(f'Documents Analyzed: {", ".join(file_names)}\n')
        doc.add_page_break()

        # Table of Contents
        doc.add_heading('Contents', level=1)
        sections = [
            '1. Executive Summary',

            '2. Financial Analysis'
        ]
        for section in sections:
            self.add_formatted_paragraph(doc, section)
        doc.add_page_break()

        # 1. Executive Summary
        doc.add_heading('1. Executive Summary', level=1)
        summary = self.generate_summary(analysis_data, anomaly_data)
        for line in summary.split('\n'):
            if line.strip():
                self.add_formatted_paragraph(doc, line)
        doc.add_page_break()

        # 2. Risk Assessment
        doc.add_heading('2. Risk Assessment', level=1)
        if 'analysis_summary' in anomaly_data:
            doc.add_heading('Risk Severity Distribution', level=2)
            severity = anomaly_data['analysis_summary']['severity_distribution']
            risk_items = [
                f"**Critical Issues:** {severity['critical']}",
                f"**High Severity Issues:** {severity['high']}",
                f"**Medium Severity Issues:** {severity['medium']}",
                f"**Low Severity Issues:** {severity['low']}"
            ]
            for item in risk_items:
                self.add_formatted_paragraph(doc, item)

        if 'anomalies' in anomaly_data:
            doc.add_heading('Detailed Risk Findings', level=2)
            for anomaly in anomaly_data['anomalies']:
                # Risk header
                self.add_formatted_paragraph(
                    doc,
                    f"**{anomaly['type'].upper()} RISK - {anomaly['severity'].upper()} Severity**",
                    'List Item'
                )
                # Description
                self.add_formatted_paragraph(
                    doc,
                    f"**Description:** {anomaly['description']}"
                )
                # Evidence
                if 'evidence' in anomaly and anomaly['evidence'].get('excerpts'):
                    self.add_formatted_paragraph(doc, "**Supporting Evidence:**")
                    for excerpt in anomaly['evidence']['excerpts']:
                        para = doc.add_paragraph(style='Quote')
                        para.add_run(excerpt)

        doc.add_page_break()

        # 3. Strategic Recommendations
        doc.add_heading('3. Strategic Recommendations', level=1)
        if 'anomalies' in anomaly_data:
            # High priority
            doc.add_heading('High Priority Actions', level=2)
            critical_items = [a for a in anomaly_data['anomalies'] if a['severity'] == 'critical']
            for item in critical_items:
                self.add_formatted_paragraph(
                    doc,
                    f"**Immediate Action Required:** Address {item['type']} risk - {item['description']}",
                    'List Item'
                )

            # Other priorities
            doc.add_heading('Additional Recommendations', level=2)
            other_items = [a for a in anomaly_data['anomalies'] if a['severity'] != 'critical']
            for item in other_items:
                self.add_formatted_paragraph(
                    doc,
                    f"**{item['severity'].title()} Priority:** {item['description']}",
                    'List Item'
                )

        doc.add_page_break()

        # 4. Financial Analysis
        doc.add_heading('4. Financial Analysis', level=1)
        if 'images' in analysis_data:
            for idx, image_data in enumerate(analysis_data['images'], 1):
                self.add_image_with_caption(
                    doc,
                    image_data['base64'],
                    f"Figure {idx}: {image_data['filename'].replace('.jpg', '')}"
                )
                doc.add_paragraph()

        # Save document
        os.makedirs('storage/reports', exist_ok=True)
        report_path = f'storage/reports/financial_report_{conversation_id}.docx'
        doc.save(report_path)
        return report_path